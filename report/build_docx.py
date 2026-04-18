"""Convert bao_cao_psntask.md → bao_cao_psntask.docx bằng python-docx.

Hỗ trợ các phần tử markdown đã dùng trong báo cáo:
- # / ## / ### headings
- **bold** + *italic* + `code` inline
- bullet list (-) và numbered list (1.)
- bảng pipe (| a | b |)
- code fence ```...```
- blockquote (>)
- horizontal rule (---)
- emoji giữ nguyên
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).parent
SRC = ROOT / "bao_cao_psntask.md"
DST = ROOT / "bao_cao_psntask.docx"


# ────────────────────────────────────────────────────────────
# Inline formatting: **bold**, *italic*, `code`
# ────────────────────────────────────────────────────────────
INLINE_RE = re.compile(
    r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)"
)


def _add_inline(paragraph, text: str, base_font="Times New Roman", size=12):
    """Thêm text có format inline (bold/italic/code) vào paragraph."""
    parts = INLINE_RE.split(text)
    for part in parts:
        if not part:
            continue
        run = paragraph.add_run()
        run.font.name = base_font
        run.font.size = Pt(size)
        if part.startswith("**") and part.endswith("**"):
            run.text = part[2:-2]
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run.text = part[1:-1]
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run.text = part[1:-1]
            run.font.name = "Consolas"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            run.text = part


def _set_cell_shading(cell, color_hex: str):
    """Tô màu nền cho ô bảng."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


# ────────────────────────────────────────────────────────────
# Block parsers
# ────────────────────────────────────────────────────────────

def _add_heading(doc, text: str, level: int):
    h = doc.add_heading(level=level)
    run = h.add_run(text)
    run.font.name = "Times New Roman"
    if level == 1:
        run.font.size = Pt(20)
        run.font.color.rgb = RGBColor(0x00, 0x58, 0xBB)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 2:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x00, 0x58, 0xBB)
    else:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2D, 0x2F, 0x33)
    run.bold = True


def _add_paragraph(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.4
    _add_inline(p, text)


def _add_bullet(doc, text: str, indent: int = 0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = 1.35
    if indent:
        p.paragraph_format.left_indent = Cm(0.6 * indent + 0.6)
    _add_inline(p, text)


def _add_numbered(doc, text: str, indent: int = 0):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.line_spacing = 1.35
    if indent:
        p.paragraph_format.left_indent = Cm(0.6 * indent + 0.6)
    _add_inline(p, text)


def _add_quote(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.35
    # bo viền trái màu xanh giả lập blockquote
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), "6C9FFF")
    borders.append(left)
    p_pr.append(borders)
    _add_inline(p, text)


def _add_code_block(doc, lines: list[str], lang: str = ""):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.space_after = Pt(8)
    # tô nền xám nhạt
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), "F0F0F6")
    p_pr.append(shd)
    text = "\n".join(lines)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x2D, 0x2F, 0x33)


def _add_table(doc, header: list[str], rows: list[list[str]]):
    table = doc.add_table(rows=1 + len(rows), cols=len(header))
    table.style = "Light Grid Accent 1"
    table.autofit = True
    # header
    for i, h in enumerate(header):
        cell = table.rows[0].cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_cell_shading(cell, "0058BB")
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # body
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = table.rows[r].cells[c]
            cell.text = ""
            p = cell.paragraphs[0]
            # value có thể chứa <br> để xuống dòng
            for k, sub in enumerate(val.split("<br>")):
                if k > 0:
                    p = cell.add_paragraph()
                _add_inline(p, sub.strip(), size=11)


def _add_hr(doc):
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "ACADB1")
    borders.append(bottom)
    p_pr.append(borders)


# ────────────────────────────────────────────────────────────
# Main parser
# ────────────────────────────────────────────────────────────

def parse_table(lines: list[str], start: int) -> tuple[int, list[str], list[list[str]]]:
    """Đọc một bảng pipe markdown bắt đầu ở `lines[start]`. Trả (next_idx, header, rows)."""
    header = [c.strip() for c in lines[start].strip().strip("|").split("|")]
    # dòng tiếp theo là separator |---|---|, bỏ qua
    i = start + 2
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        # bù cột nếu thiếu
        while len(row) < len(header):
            row.append("")
        rows.append(row[: len(header)])
        i += 1
    return i, header, rows


def parse_md(path: Path) -> Document:
    doc = Document()

    # default style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    # margin
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    raw = path.read_text(encoding="utf-8")
    lines = raw.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # blank
        if not stripped:
            i += 1
            continue

        # horizontal rule
        if stripped == "---":
            _add_hr(doc)
            i += 1
            continue

        # headings
        if stripped.startswith("# "):
            _add_heading(doc, stripped[2:].strip(), 1)
            i += 1
            continue
        if stripped.startswith("## "):
            _add_heading(doc, stripped[3:].strip(), 2)
            i += 1
            continue
        if stripped.startswith("### "):
            _add_heading(doc, stripped[4:].strip(), 3)
            i += 1
            continue

        # code fence
        if stripped.startswith("```"):
            lang = stripped[3:].strip()
            i += 1
            buf: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # bỏ qua ```
            _add_code_block(doc, buf, lang)
            continue

        # blockquote
        if stripped.startswith("> "):
            buf = [stripped[2:]]
            i += 1
            while i < len(lines) and lines[i].strip().startswith("> "):
                buf.append(lines[i].strip()[2:])
                i += 1
            for ln in buf:
                _add_quote(doc, ln)
            continue

        # table
        if stripped.startswith("|") and i + 1 < len(lines) \
                and re.match(r"^\|[\s:-]+\|", lines[i + 1].strip()):
            i, header, rows = parse_table(lines, i)
            _add_table(doc, header, rows)
            continue

        # bullet list
        m = re.match(r"^(\s*)-\s+(.*)$", line)
        if m:
            indent = len(m.group(1)) // 2
            _add_bullet(doc, m.group(2), indent)
            i += 1
            continue

        # numbered list
        m = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if m:
            indent = len(m.group(1)) // 2
            _add_numbered(doc, m.group(2), indent)
            i += 1
            continue

        # plain paragraph
        _add_paragraph(doc, stripped)
        i += 1

    return doc


def main():
    if not SRC.exists():
        raise SystemExit(f"Không thấy {SRC}")
    doc = parse_md(SRC)
    doc.save(DST)
    print(f"✓ Đã sinh: {DST}  ({DST.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
